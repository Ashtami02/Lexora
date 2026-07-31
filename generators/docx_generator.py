import os

from docx import Document


class DOCXGenerator:

    def generate(
        self,
        output_path: str,
        title: str,
        document_text: str,
    ):

        os.makedirs(
            os.path.dirname(output_path),
            exist_ok=True,
        )

        document = Document()

        document.add_heading(title, level=1)

        document.add_paragraph(document_text)

        document.save(output_path)

        filename = os.path.basename(output_path)

        return {
            "path": output_path,
            "url": f"/generated_documents/{filename}",
        }


docx_generator = DOCXGenerator()